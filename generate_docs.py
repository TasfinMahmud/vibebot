"""Generate VibeBot Complete User Guide as a Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style Setup ─────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Segoe UI"
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x34, 0x36)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f"Heading {level}"]
    heading_style.font.name = "Segoe UI"
    heading_style.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    if level == 1:
        heading_style.font.size = Pt(22)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
    else:
        heading_style.font.size = Pt(13)
        heading_style.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)

# ── Helper Functions ────────────────────────────────
def add_colored_text(paragraph, text, color=None, bold=False, size=None):
    run = paragraph.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    return run

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Add shading
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    shading.set(qn("w:val"), "clear")
    run._element.get_or_add_rPr().append(shading)

def add_tip(doc, text, label="💡 TIP"):
    p = doc.add_paragraph()
    add_colored_text(p, f"{label}: ", RGBColor(0x05, 0x96, 0x69), bold=True, size=Pt(10))
    add_colored_text(p, text, RGBColor(0x05, 0x96, 0x69), size=Pt(10))

def add_warning(doc, text, label="⚠️ WARNING"):
    p = doc.add_paragraph()
    add_colored_text(p, f"{label}: ", RGBColor(0xD9, 0x77, 0x06), bold=True, size=Pt(10))
    add_colored_text(p, text, RGBColor(0xD9, 0x77, 0x06), size=Pt(10))

# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(title_p, "🤖 VibeBot", RGBColor(0x7C, 0x3A, 0xED), bold=True, size=Pt(36))

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(subtitle_p, "Complete User Guide & Documentation", RGBColor(0x6B, 0x72, 0x80), size=Pt(18))

doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(desc_p, "Your server's all-in-one fun companion!\n", RGBColor(0x4B, 0x55, 0x63), size=Pt(12))
add_colored_text(desc_p, "Announcements • Insights • Auto-Replies • Games • Polls • Proxy Messaging", RGBColor(0x6B, 0x72, 0x80), size=Pt(11))

for _ in range(4):
    doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(meta_p, "Version 1.0  •  April 2026  •  Bot Tag: VibeBot#3323", RGBColor(0x9C, 0xA3, 0xAF), size=Pt(10))

doc.add_page_break()

# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Quick Start",
    "2. Post as Bot (Proxy Messaging)",
    "3. Deadline Reminders",
    "4. Announcement Commands",
    "5. Server Insights",
    "6. Fun Commands",
    "7. Utility Commands",
    "8. Auto-Reply System",
    "9. Mention Responder",
    "10. Welcome System",
    "11. Command Cheat Sheet",
    "12. Customization Guide",
    "13. Troubleshooting",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. QUICK START
# ═══════════════════════════════════════════════════
doc.add_heading("1. Quick Start", level=1)
doc.add_paragraph("Get VibeBot up and running in 3 simple steps:")
make_table(doc,
    ["Step", "Command", "What It Does"],
    [
        ["1", "npm run deploy", "Registers all slash commands with Discord"],
        ["2", "npm start", "Starts the bot"],
        ["3", "Type /help in Discord", "Shows all available commands"],
    ])
doc.add_paragraph()
add_tip(doc, "Slash commands may take up to 1 hour to appear globally after first registration. They usually show up within a few minutes.")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 2. POST AS BOT (PROXY MESSAGING)
# ═══════════════════════════════════════════════════
doc.add_heading("2. Post as Bot (Proxy Messaging)", level=1)
p = doc.add_paragraph()
add_colored_text(p, "🔒 Owner Only", RGBColor(0xEF, 0x44, 0x44), bold=True, size=Pt(11))
add_colored_text(p, " — Only the server owner can use this command.", size=Pt(11))

doc.add_paragraph("Send messages through the bot so it appears as if the bot is posting. Your slash command is hidden (ephemeral) — only the resulting message is visible to everyone.")

doc.add_heading("/post — Send a Categorized Message", level=2)
make_table(doc,
    ["Parameter", "Required", "Description", "Example"],
    [
        ["category", "✅ Yes", "Message style (Announcement, Insight, Info, Reminder)", "📢 Announcement"],
        ["message", "✅ Yes", "The message content", "Midterm is on April 15th!"],
        ["title", "❌ No", "Optional custom heading", "Important Update"],
        ["channel", "❌ No", "Target channel (defaults to current)", "#announcements"],
        ["ping_everyone", "❌ No", "Ping @everyone with this message?", "True"],
    ])

doc.add_paragraph()
doc.add_heading("Category Styles", level=3)
make_table(doc,
    ["Category", "Emoji", "Color", "Best For"],
    [
        ["Announcement", "📢", "Purple", "Official server announcements"],
        ["Insight", "💡", "Amber/Yellow", "Tips, study insights, shared knowledge"],
        ["Info", "ℹ️", "Blue", "General information, updates"],
        ["Reminder", "🔔", "Red", "Deadlines, meeting reminders"],
    ])

doc.add_paragraph()
doc.add_heading("Usage Examples", level=3)
add_code_block(doc, "/post category:📢 Announcement message:Midterm is on April 15th!")
add_code_block(doc, "/post category:💡 Insight message:Pro tip: Use Anki flashcards for better retention! title:Study Hack")
add_code_block(doc, '/post category:🔔 Reminder message:Submit CSE301 assignment by tonight! ping_everyone:True')
add_code_block(doc, "/post category:ℹ️ Info message:Office hours moved to 3 PM channel:#announcements")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 3. DEADLINE REMINDERS
# ═══════════════════════════════════════════════════
doc.add_heading("3. Deadline Reminders", level=1)
p = doc.add_paragraph()
add_colored_text(p, "🔒 Owner Only", RGBColor(0xEF, 0x44, 0x44), bold=True, size=Pt(11))

doc.add_heading("/remind — Schedule a Deadline Reminder", level=2)
doc.add_paragraph("Schedule a reminder that automatically fires after a set time, pinging @everyone with a big red alert embed.")

make_table(doc,
    ["Parameter", "Required", "Description", "Example"],
    [
        ["title", "✅ Yes", "What's the reminder about?", "Assignment Due!"],
        ["message", "✅ Yes", "Reminder details", "CSE301 assignment deadline in 30 mins!"],
        ["minutes", "✅ Yes", "Minutes from now to fire", "30"],
        ["channel", "❌ No", "Channel to post in", "#general"],
    ])

doc.add_paragraph()
doc.add_heading("Usage Example", level=3)
add_code_block(doc, "/remind title:Assignment Due! message:CSE301 assignment deadline! Submit now! minutes:30 channel:#general")
doc.add_paragraph("After 30 minutes, the bot will automatically post:")
add_code_block(doc, "@everyone 🔔 DEADLINE REMINDER\n┌──────────────────────────────────────────┐\n│ 🚨 REMINDER: Assignment Due!             │\n│                                          │\n│ CSE301 assignment deadline! Submit now!   │\n│                                          │\n│ ⏰ This is your scheduled reminder!       │\n│ 📅 Deadline alert — please take action!   │\n└──────────────────────────────────────────┘")

add_warning(doc, "Scheduled reminders are stored in memory. If the bot restarts before the timer fires, the reminder will be lost.")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 4. ANNOUNCEMENT COMMANDS
# ═══════════════════════════════════════════════════
doc.add_heading("4. Announcement Commands", level=1)
doc.add_paragraph("These commands require the Manage Messages permission.")

doc.add_heading("/announce — Post a Fancy Announcement", level=2)
make_table(doc,
    ["Parameter", "Required", "Description", "Example"],
    [
        ["title", "✅ Yes", "Announcement heading", "Server Maintenance"],
        ["message", "✅ Yes", "Body text", "Server will be down for 2 hours tonight"],
        ["channel", "❌ No", "Target channel (defaults to current)", "#announcements"],
        ["color", "❌ No", "Custom hex color for embed", "#ff6b6b"],
        ["image", "❌ No", "URL of an image to attach", "https://example.com/img.png"],
    ])

doc.add_paragraph()
doc.add_heading("Usage Example", level=3)
add_code_block(doc, "/announce title:🎮 Game Night! message:Join us this Friday at 9 PM for Among Us! channel:#announcements color:#7c3aed")

doc.add_heading("/schedule — Schedule an Announcement", level=2)
make_table(doc,
    ["Parameter", "Required", "Description", "Example"],
    [
        ["title", "✅ Yes", "Announcement heading", "Reminder!"],
        ["message", "✅ Yes", "Body text", "Meeting starts in 5 minutes"],
        ["minutes", "✅ Yes", "Delay in minutes", "30"],
        ["channel", "❌ No", "Target channel", "#general"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 5. SERVER INSIGHTS
# ═══════════════════════════════════════════════════
doc.add_heading("5. Server Insights", level=1)

doc.add_heading("/insight — Full Server Statistics", level=2)
doc.add_paragraph("Displays a beautiful embed with comprehensive server statistics. No parameters needed — just type /insight!")
make_table(doc,
    ["Stat", "Description"],
    [
        ["👥 Members", "Total member count"],
        ["🟢 Online", "Currently online members"],
        ["🤖 Bots", "Number of bots in the server"],
        ["💬 Text Channels", "Count of text channels"],
        ["🔊 Voice Channels", "Count of voice channels"],
        ["🎭 Roles", "Number of roles"],
        ["😄 Emojis", "Number of custom emojis"],
        ["💎 Boosts", "Nitro boost count & tier"],
        ["📅 Created", "When the server was created"],
    ])

doc.add_paragraph()
doc.add_heading("/userinfo — User Lookup", level=2)
doc.add_paragraph("Get detailed info about any server member. Shows username, ID, account creation date, server join date, roles, and avatar.")
add_code_block(doc, "/userinfo user:@SomeFriend")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 6. FUN COMMANDS
# ═══════════════════════════════════════════════════
doc.add_heading("6. Fun Commands", level=1)

commands = [
    ("/joke", "😂 Random Joke", "Tells a random joke with the punchline hidden behind a spoiler tag — click to reveal!", "/joke", "12 curated jokes"),
    ("/8ball", "🔮 Magic 8-Ball", "Ask any yes/no question and get a mystical answer!", "/8ball question:Will I pass my exam?", "20 unique responses"),
    ("/funfact", "🧠 Random Fun Fact", "Get a mind-blowing random fact about science, nature, or tech!", "/funfact", "18 unique facts"),
    ("/quote", "💬 Inspirational Quote", "Get motivation from famous people like Steve Jobs, Einstein, and Churchill!", "/quote", "14 curated quotes"),
    ("/vibe", "✨ Vibe Check", "Check your random vibe level for the day! Shows a percentage bar from 10% to 100%.", "/vibe", "10 vibe levels"),
    ("/coinflip", "🪙 Coin Flip", "Flip a coin — Heads or Tails!", "/coinflip", "50/50 odds"),
    ("/roll", "🎲 Dice Roll", "Roll a dice with custom sides (default d6).", "/roll sides:20", "Any number of sides"),
]

for cmd, title, desc, example, pool in commands:
    doc.add_heading(f"{cmd} — {title}", level=2)
    doc.add_paragraph(desc)
    add_code_block(doc, example)
    p = doc.add_paragraph()
    add_colored_text(p, f"Pool: {pool}", RGBColor(0x6B, 0x72, 0x80), size=Pt(9))

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 7. UTILITY COMMANDS
# ═══════════════════════════════════════════════════
doc.add_heading("7. Utility Commands", level=1)

doc.add_heading("/poll — Create a Poll", level=2)
doc.add_paragraph("Create an interactive poll with emoji reactions. Supports 2-4 options.")
make_table(doc,
    ["Parameter", "Required", "Description"],
    [
        ["question", "✅ Yes", "The poll question"],
        ["option1", "✅ Yes", "First option"],
        ["option2", "✅ Yes", "Second option"],
        ["option3", "❌ No", "Third option"],
        ["option4", "❌ No", "Fourth option"],
    ])
doc.add_paragraph()
add_code_block(doc, "/poll question:Best programming language? option1:Python option2:JavaScript option3:Java option4:C++")
doc.add_paragraph("Members vote by clicking the emoji reactions (1️⃣ 2️⃣ 3️⃣ 4️⃣) that are automatically added.")

doc.add_heading("/avatar — Full-Size Avatar", level=2)
doc.add_paragraph("Shows a user's full 1024px avatar with a 'Open in Browser' button for the 4096px version.")
add_code_block(doc, "/avatar user:@Friend")

doc.add_heading("/help — Command List", level=2)
doc.add_paragraph("Shows a categorized list of all available commands with descriptions.")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 8. AUTO-REPLY SYSTEM
# ═══════════════════════════════════════════════════
doc.add_heading("8. Auto-Reply System", level=1)
doc.add_paragraph("VibeBot automatically responds to certain keywords in chat messages. No commands needed — just chat naturally!")

make_table(doc,
    ["Trigger Words", "Example Responses"],
    [
        ["hello, hi", 'Hey there! 👋  •  Yo! What\'s good? 🎉'],
        ["good morning", "Good morning! ☀️ Rise and grind!"],
        ["good night", "Sweet dreams! 🌙✨"],
        ["thanks, thank you", "No prob! 😄  •  Anytime! 💪"],
        ["help", "Need help? Try /help to see what I can do! 🤖"],
        ["lol", "😂😂😂  •  LMAOOO 💀"],
        ["bruh", "Bruh moment certified 💀"],
        ["sad", "Aww don't be sad! Here's a virtual hug 🤗"],
        ["exam", "Good luck on your exam! 📚🍀"],
        ["assignment", "Assignment grind never stops 📝"],
        ["bored", "Bored? Try /joke or /8ball! 🎲"],
    ])

doc.add_paragraph()
doc.add_heading("ALL CAPS Detection (Easter Egg)", level=2)
doc.add_paragraph("If someone types a message longer than 10 characters in ALL CAPS, the bot calls them out with fun responses like:")
doc.add_paragraph('"Whoa, easy on the caps lock! 😅"', style="List Bullet")
doc.add_paragraph('"I CAN HEAR YOU LOUD AND CLEAR 📢"', style="List Bullet")
doc.add_paragraph('"Sir/Ma\'am, this is a Discord server 🤣"', style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 9. MENTION RESPONDER
# ═══════════════════════════════════════════════════
doc.add_heading("9. Mention Responder", level=1)
doc.add_paragraph("When someone @mentions another user in a message, VibeBot automatically replies to let them know the mentioned person will respond when they see the message.")
doc.add_paragraph()
doc.add_heading("How It Works", level=2)
doc.add_paragraph("Someone types: @John how are you?")
doc.add_paragraph("VibeBot replies:")
add_code_block(doc, "👋 Hey! John has been notified. He/She will respond once he/she sees the message! 📩")
doc.add_paragraph()
doc.add_paragraph("Supports multiple mentions and gives grammatically correct responses (singular/plural).")
add_tip(doc, "The bot ignores self-mentions and bot mentions — it only responds when a real user is mentioned by someone else.")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 10. WELCOME SYSTEM
# ═══════════════════════════════════════════════════
doc.add_heading("10. Welcome System", level=1)
doc.add_paragraph("When a new member joins the server, VibeBot automatically sends a beautiful welcome embed in the system/general channel showing:")
doc.add_paragraph("Personalized welcome message", style="List Bullet")
doc.add_paragraph("Member's avatar", style="List Bullet")
doc.add_paragraph("Member count (You're member #42!)", style="List Bullet")
doc.add_paragraph("Account creation date", style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 11. COMMAND CHEAT SHEET
# ═══════════════════════════════════════════════════
doc.add_heading("11. Command Cheat Sheet", level=1)

make_table(doc,
    ["Command", "Category", "Description", "Who Can Use"],
    [
        ["/post", "✉️ Proxy", "Send a message as the bot (4 categories)", "🔒 Owner only"],
        ["/remind", "✉️ Proxy", "Schedule a deadline reminder with @everyone", "🔒 Owner only"],
        ["/announce", "📢 Announce", "Post a fancy announcement embed", "Manage Messages"],
        ["/schedule", "📢 Announce", "Schedule a timed announcement", "Manage Messages"],
        ["/insight", "📊 Insights", "Full server stats and info", "Everyone"],
        ["/userinfo", "📊 Insights", "Look up a user's details", "Everyone"],
        ["/joke", "🎮 Fun", "Random joke with spoiler punchline", "Everyone"],
        ["/8ball", "🎮 Fun", "Ask the magic 8-ball", "Everyone"],
        ["/funfact", "🎮 Fun", "Random fun fact", "Everyone"],
        ["/quote", "🎮 Fun", "Inspirational quote", "Everyone"],
        ["/vibe", "🎮 Fun", "Check your vibe level", "Everyone"],
        ["/coinflip", "🎮 Fun", "Flip a coin", "Everyone"],
        ["/roll", "🎮 Fun", "Roll a dice (custom sides)", "Everyone"],
        ["/poll", "🛠️ Utility", "Create interactive poll", "Everyone"],
        ["/avatar", "🛠️ Utility", "Get full-size avatar", "Everyone"],
        ["/help", "🛠️ Utility", "Show all commands", "Everyone"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 12. CUSTOMIZATION GUIDE
# ═══════════════════════════════════════════════════
doc.add_heading("12. Customization Guide", level=1)
doc.add_paragraph("All bot content is easily customizable in config.js. After editing, just restart the bot (Ctrl+C then npm start) — no need to re-deploy commands.")

doc.add_heading("Add New Auto-Replies", level=2)
add_code_block(doc, '// In config.js → autoReplies object\nautoReplies: {\n  "your trigger": ["Response 1", "Response 2", "Response 3"],\n}')

doc.add_heading("Add New Jokes", level=2)
add_code_block(doc, '// In config.js → jokes array\njokes: [\n  { setup: "Your joke setup?", punchline: "The punchline! 😂" },\n]')

doc.add_heading("Change Bot Colors", level=2)
add_code_block(doc, "// In config.js → accentColors\naccentColors: {\n  success: 0x10b981,     // Green\n  warning: 0xf59e0b,     // Yellow\n  error: 0xef4444,       // Red\n  info: 0x3b82f6,        // Blue\n  fun: 0xec4899,         // Pink\n  announcement: 0x7c3aed, // Purple\n}")

doc.add_heading("Change Bot Name", level=2)
add_code_block(doc, 'botName: "YourBotName",')

add_tip(doc, "After editing config.js, just restart the bot — no need to re-deploy commands.")

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 13. TROUBLESHOOTING
# ═══════════════════════════════════════════════════
doc.add_heading("13. Troubleshooting", level=1)

make_table(doc,
    ["Problem", "Solution"],
    [
        ["Bot is offline", "Run npm start in the terminal"],
        ["Slash commands don't appear", "Run npm run deploy and wait up to 1 hour"],
        ['"Missing permissions" error', "Make sure you have Manage Messages permission"],
        ["Auto-replies not working", "Ensure Message Content Intent is enabled in Developer Portal"],
        ["Welcome messages not showing", "Ensure Server Members Intent is enabled"],
        ["Bot can't see online count", "Ensure Presence Intent is enabled"],
        ["/post says 'only server owner'", "This command is restricted to the server owner only"],
        ["Bot goes offline when PC sleeps", "Deploy to a cloud service (Railway, Render) for 24/7 uptime"],
    ])

doc.add_paragraph()
add_warning(doc, "Never share your bot token! If it's leaked, go to Discord Developer Portal → Bot → Reset Token immediately.")

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ──────────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(footer_p, "━" * 40, RGBColor(0xD1, 0xD5, 0xDB))
footer_p2 = doc.add_paragraph()
footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(footer_p2, "VibeBot v1.0 • Built with discord.js • April 2026", RGBColor(0x9C, 0xA3, 0xAF), size=Pt(9))
footer_p3 = doc.add_paragraph()
footer_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_colored_text(footer_p3, "Making servers fun since 2026 ✨", RGBColor(0x7C, 0x3A, 0xED), size=Pt(10))

# ── Save ────────────────────────────────────────────
output_path = r"e:\301\discord-bot\VibeBot_User_Guide.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
